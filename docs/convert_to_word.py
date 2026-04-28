"""
Conversion Markdown → Word (.docx)
Dépendance : pip install python-docx
Usage     : python convert_to_word.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Chemins ────────────────────────────────────────────────────────────────
INPUT_MD  = Path(__file__).parent / "rapport_final_projet.md"
OUTPUT_DOC = Path(__file__).parent / "rapport_final_projet.docx"


# ── Utilitaires de style ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Colore le fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document):
    """Ajoute un séparateur horizontal."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def apply_inline_styles(paragraph, text: str):
    """
    Parse le texte inline et ajoute des runs stylisés :
    **gras**, *italique*, `code inline`, ~~barré~~
    """
    # Regex qui capture les marqueurs inline dans l'ordre de priorité
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # gras+italique
        r"|(\*\*(.+?)\*\*)"       # gras
        r"|(\*(.+?)\*)"           # italique
        r"|(`(.+?)`)"             # code inline
        r"|(~~(.+?)~~)"           # barré
        r"|(\[([^\]]+)\]\([^\)]+\))"  # lien [texte](url) → texte seulement
    )
    last = 0
    for m in pattern.finditer(text):
        # Texte brut avant le marqueur
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])

        if m.group(1):   # gras+italique
            run = paragraph.add_run(m.group(2))
            run.bold = True; run.italic = True
        elif m.group(3): # gras
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group(5): # italique
            run = paragraph.add_run(m.group(6))
            run.italic = True
        elif m.group(7): # code inline
            run = paragraph.add_run(m.group(8))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        elif m.group(9): # barré
            run = paragraph.add_run(m.group(10))
            run.font.strike = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        elif m.group(11): # lien → texte uniquement
            run = paragraph.add_run(m.group(12))
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

        last = m.end()

    # Reste du texte
    if last < len(text):
        paragraph.add_run(text[last:])


def clean_inline(text: str) -> str:
    """Retire les marqueurs markdown pour un texte brut (ex: titres)."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*",     r"\1", text)
    text = re.sub(r"\*(.+?)\*",          r"\1", text)
    text = re.sub(r"`(.+?)`",            r"\1", text)
    text = re.sub(r"~~(.+?)~~",          r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text


# ── Parser principal ─────────────────────────────────────────────────────────

def parse_markdown_to_docx(md_text: str, doc: Document):
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Bloc de code ────────────────────────────────────────────────────
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph(style="No Spacing")
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # Fond gris clair via shading sur le paragraphe
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F0F0F0")
            pPr.append(shd)
            i += 1
            continue

        # ── Tableau ─────────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Filtrer la ligne de séparation (|----|)
            data_rows = [
                r for r in table_lines
                if not re.match(r"^\|[\s\-\:\|]+\|$", r)
            ]
            if not data_rows:
                continue

            def split_row(row):
                return [c.strip() for c in row.strip().strip("|").split("|")]

            headers = split_row(data_rows[0])
            n_cols  = len(headers)

            table = doc.add_table(rows=1, cols=n_cols)
            table.style = "Table Grid"

            # En-têtes
            hdr_cells = table.rows[0].cells
            for j, h in enumerate(headers):
                hdr_cells[j].text = clean_inline(h)
                hdr_cells[j].paragraphs[0].runs[0].bold = True
                set_cell_bg(hdr_cells[j], "D9E1F2")

            # Données
            for row_line in data_rows[1:]:
                cells_data = split_row(row_line)
                row_cells  = table.add_row().cells
                for j, cell_text in enumerate(cells_data):
                    if j < n_cols:
                        p = row_cells[j].paragraphs[0]
                        apply_inline_styles(p, cell_text)

            doc.add_paragraph()  # espace après le tableau
            continue

        # ── Séparateur horizontal ────────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Titres ──────────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = clean_inline(heading_match.group(2))
            doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Blockquote ──────────────────────────────────────────────────────
        if line.startswith("> "):
            content = line[2:]
            p = doc.add_paragraph(style="No Spacing")
            p.paragraph_format.left_indent  = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "4472C4")
            pBdr.append(left)
            pPr.append(pBdr)
            apply_inline_styles(p, content)
            i += 1
            continue

        # ── Liste à puces ────────────────────────────────────────────────────
        if re.match(r"^(\s*)[-*+]\s+(.+)", line):
            m = re.match(r"^(\s*)[-*+]\s+(.+)", line)
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            apply_inline_styles(p, m.group(2))
            i += 1
            continue

        # ── Liste numérotée ──────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+(.+)", line):
            m = re.match(r"^\d+\.\s+(.+)", line)
            p = doc.add_paragraph(style="List Number")
            apply_inline_styles(p, m.group(1))
            i += 1
            continue

        # ── Métadonnées de début (gras/valeur) ───────────────────────────────
        if re.match(r"^\*\*[^*]+\*\*\s*:", line):
            p = doc.add_paragraph(style="No Spacing")
            apply_inline_styles(p, line)
            i += 1
            continue

        # ── Ligne vide ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Paragraphe normal ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        apply_inline_styles(p, line)
        i += 1


# ── Configuration du document ─────────────────────────────────────────────────

def configure_document(doc: Document):
    """Marges, police par défaut, styles de titres."""
    from docx.shared import Cm as cm

    # Marges
    for section in doc.sections:
        section.top_margin    = cm(2.5)
        section.bottom_margin = cm(2.5)
        section.left_margin   = cm(2.5)
        section.right_margin  = cm(2.5)

    # Police par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Couleurs des titres
    title_colors = {
        "Heading 1": (0x26, 0x3C, 0x78),  # Bleu foncé
        "Heading 2": (0x2E, 0x74, 0xB5),  # Bleu moyen
        "Heading 3": (0x2E, 0x74, 0xB5),  # Bleu moyen (plus petit)
    }
    for style_name, (r, g, b) in title_colors.items():
        try:
            s = doc.styles[style_name]
            s.font.color.rgb = RGBColor(r, g, b)
        except KeyError:
            pass


# ── Point d'entrée ────────────────────────────────────────────────────────────

def main():
    print(f"Lecture : {INPUT_MD}")
    md_text = INPUT_MD.read_text(encoding="utf-8")

    doc = Document()
    configure_document(doc)
    parse_markdown_to_docx(md_text, doc)

    doc.save(OUTPUT_DOC)
    print(f"Document Word créé : {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
